# -*- coding: utf-8 -*-
"""
RTDB 매뉴얼 공통 생성기 (정본 — 지우지 말 것)
==============================================

왜 있는가
---------
매뉴얼 생성기를 임시로 만들었다가 버리는 일이 반복되면서, 나중에 문서를 고칠 때마다
docx 를 손으로 뜯어고쳐야 했다. 그 과정에서 서식이 조금씩 어긋난다.
그래서 이 파일을 **저장소 정본**으로 두고 모든 RTDB 매뉴얼이 이것만 쓰도록 한다.

서식 기준 (기존 매뉴얼 6권에서 실측한 값 — 임의로 바꾸지 말 것)
----------------------------------------------------------------
  용지      A4, 여백 좌우 2.2cm / 상하 2.0cm
  본문      맑은 고딕 10.5pt
  제목 1    맑은 고딕 17pt   굵게  #005A9E
  제목 2    맑은 고딕 13.5pt 굵게  #005A9E
  제목 3    맑은 고딕 11.5pt 굵게  #333333
  표        Table Grid · 모든 칸 세로 가운데 · 제목 행은 가로·세로 가운데
  그림      폭 12.5cm 기본 · 높이 9.0cm 상한 · 원본의 1.6배 넘게 확대 안 함
  그림 설명 9pt #777777, 가운데
  표지      제목 26pt #005A9E / 부제 13pt #555555 / 정보 10.5pt #666666

쪽 넘김 규칙 (2026-08-13 추가 — 전 제품 공통)
--------------------------------------------
쪽 마지막에 한두 줄만 남고 본체가 다음 장으로 넘어가는 조판을 막는다.
  · 그림 : 그림과 그림 설명을 붙인다 + 높이 상한 7.5cm 로 크기 자체를 억제
  · 표   : 행 분할 금지(`cantSplit`) + 쪽 넘어가면 제목 행 반복(`tblHeader`)
  · 상자 : 통째로 넘긴다(`keep_together`)
  · 본문 : `widow_control`
  · 제목 : 뒤 문단과 붙인다(`keep_with_next`) — 제목만 쪽 끝에 남지 않게

⚠ 묶음을 «연쇄»시키면 역효과가 난다. 표·그림 앞 문단까지 붙였더니
  제목→도입문→그림→설명 넷이 한 덩어리가 되어 통째로 밀리고, 앞 쪽에 10~14cm
  빈 공간이 남았다(실측). 묶음은 갈라지면 안 되는 짝에만 건다.
  **크기를 줄이는 것이 규칙보다 먼저다.**

쓰는 법
-------
    from manual_kit import ManualDoc
    m = ManualDoc("RTDB 이상감지", "부제", version="1.0", date="2026-07-30")
    m.toc([("1. 개요", []), ("2. 기법", ["2-1. EWMA"])])
    m.h1("1. 개요"); m.p("본문...")
    m.table(["열1","열2"], [["a","b"]])
    m.figure("images/fig.png", "[그림 1] 설명")
    m.save("경로.docx")

PDF 변환은 tools/docx2pdf.ps1 (Word COM) 을 쓴다.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "맑은 고딕"
NAVY = "005A9E"
DARK = "333333"
GRAY = "555555"
INFO = "666666"
CAPTION = "777777"
BOXBG = "F2F7FB"      # 참고 상자 배경(연한 하늘)
BOXBG_WARN = "FDF3E7"  # 주의 상자 배경(연한 주황)

# ── 그림 크기 규칙 (전 제품 매뉴얼 공통) ───────────────────────
# 폭 하나만 지정하면 큰 캡처는 쪽을 잡아먹고 작은 조각 캡처는 억지로 늘어나
# 뭉개진다. 그래서 «폭·높이·확대배율» 세 가지로 함께 묶는다.
PX_PER_CM = 37.795     # 96 DPI 기준 — 원본을 화면과 같은 크기로 볼 때의 환산값
FIG_W_CM = 12.5        # 기본 폭
FIG_MAX_H_CM = 7.5     # 높이 상한 — 본문 높이(25.7cm)의 3할. 이보다 크면 쪽에 못 들어가
                       # 통째로 밀리면서 앞 쪽에 큰 빈 공간이 남는다(실측으로 정한 값)
FIG_MAX_UPSCALE = 1.6  # 확대 상한 — 작은 캡처를 이 배율 이상으로 늘리지 않는다


def _apply_font(obj, name=FONT, size=None, bold=None, italic=None, color=None):
    """run 또는 style 에 폰트를 건다. 한글은 eastAsia 를 같이 지정해야 적용된다."""
    f = obj.font
    f.name = name
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = RGBColor.from_string(color)
    rpr = obj.element.get_or_add_rPr() if hasattr(obj, "element") else obj._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), name)


def _shade(cell_or_par, hex_color):
    """셀/문단 배경색.

    ⚠ 셀은 w:tc 가 아니라 w:tcPr 에 넣어야 한다. w:tc 에 바로 붙이면 Word 가 조용히
      무시하고, 머리행 흰 글자가 흰 배경 위에 찍혀 글자가 사라진 것처럼 보인다.
    """
    el = (cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, "_tc")
          else cell_or_par._p.get_or_add_pPr())
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    el.append(shd)


def _border(par, color=NAVY, size=6, sides=("left",)):
    """문단 테두리(참고 상자 좌측 굵은 선)."""
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for s in sides:
        e = OxmlElement("w:" + s)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), color)
        bdr.append(e)
    pPr.append(bdr)


def _split_bold(text):
    """`**굵게**` 를 [(글자, 굵게), …] 로 나눈다. 표시가 없으면 통째로 한 덩어리."""
    parts, buf, bold = [], [], False
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            if buf:
                parts.append(("".join(buf), bold))
                buf = []
            bold = not bold
            i += 2
            continue
        buf.append(text[i])
        i += 1
    if buf:
        parts.append(("".join(buf), bold))
    return parts or [("", False)]


class ManualDoc:
    def __init__(self, title, subtitle, version="1.0", date="", company="바이트론 이앤에스"):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._footer_pagenum()
        self._cover(title, subtitle, version, date, company)

    # ── 기본 설정 ───────────────────────────────────────────────
    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2.0)

    def _setup_styles(self):
        st = self.doc.styles
        _apply_font(st["Normal"], size=10.5)
        st["Normal"].paragraph_format.space_after = Pt(6)
        st["Normal"].paragraph_format.line_spacing = 1.35
        # 문단이 쪽 경계에서 한 줄만 떼어져 남지 않게 한다(고아·미망인 줄 방지).
        st["Normal"].paragraph_format.widow_control = True

        for name, size, color, before, after in (
            ("Heading 1", 17, NAVY, 20, 10),
            ("Heading 2", 13.5, NAVY, 14, 6),
            ("Heading 3", 11.5, DARK, 10, 4),
        ):
            s = st[name]
            _apply_font(s, size=size, bold=True, color=color)
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.keep_with_next = True

    def _footer_pagenum(self):
        p = self.doc.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        _apply_font(run, size=9, color=CAPTION)
        for el, attr in (("w:fldChar", "begin"), ("w:instrText", None), ("w:fldChar", "end")):
            e = OxmlElement(el)
            if el == "w:fldChar":
                e.set(qn("w:fldCharType"), attr)
            else:
                e.set(qn("xml:space"), "preserve")
                e.text = " PAGE "
            run._r.append(e)

    def _cover(self, title, subtitle, version, date, company):
        for _ in range(4):
            self.doc.add_paragraph()
        self._centered(title, 26, NAVY, bold=True)
        self._centered(subtitle, 13, GRAY)
        for _ in range(8):
            self.doc.add_paragraph()
        self._centered(f"문서 버전  {version}", 10.5, INFO)
        if date:
            self._centered(f"작성일  {date}", 10.5, INFO)
        self._centered(company, 10.5, INFO)
        self.page_break()

    def _centered(self, text, size, color, bold=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        _apply_font(r, size=size, bold=bold, color=color)
        return p

    # ── 본문 요소 ───────────────────────────────────────────────
    def page_break(self):
        self.doc.add_page_break()

    def h1(self, text, new_page=False):
        """기존 RTDB 매뉴얼과 동일하게 장은 이어서 흐른다(장마다 쪽을 끊으면
        장 끝의 한두 줄만 남는 빈 쪽이 계속 생긴다). 끊고 싶을 때만 new_page=True."""
        if new_page:
            self.page_break()
        return self.doc.add_heading(text, level=1)

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def p(self, text="", bold=False, size=10.5, color=None, indent=0, align=None):
        par = self.doc.add_paragraph()
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        if align == "center":
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            # ⚠ 문단 안의 `**굵게**` 를 해석한다. 이게 없으면 별표가 그대로 찍힌다
            #   (2026-08-06 전원품질 매뉴얼에서 실제로 겪음).
            #   표시가 없는 문단은 종전과 똑같이 한 덩어리로 처리된다.
            for txt, b in _split_bold(text):
                r = par.add_run(txt)
                _apply_font(r, size=size, bold=(bold or b), color=color)
        return par

    def rich(self, parts, indent=0):
        """[(텍스트, 굵게), ...] 로 한 문단 안에서 굵기를 섞는다."""
        par = self.doc.add_paragraph()
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        for text, bold in parts:
            r = par.add_run(text)
            _apply_font(r, size=10.5, bold=bold)
        return par

    def bullet(self, text, level=0):
        par = self.doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
        par.paragraph_format.space_after = Pt(3)
        mark = par.add_run("· " if level == 0 else "- ")
        _apply_font(mark, size=10.5)
        # p()·note()·table() 과 같이 `**굵게**` 를 해석한다.
        # 이게 없으면 별표가 그대로 찍힌다.
        for txt, b in _split_bold(text):
            r = par.add_run(txt)
            _apply_font(r, size=10.5, bold=b)
        return par

    def steps(self, items):
        """절차 목록. Word 자동 번호는 문서 전체에서 이어 세는 사고가 있어 직접 번호를 찍는다."""
        for i, t in enumerate(items, 1):
            par = self.doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.6)
            par.paragraph_format.space_after = Pt(3)
            r = par.add_run(f"{i}. ")
            _apply_font(r, size=10.5, bold=True)
            # `**굵게**` 해석 — 없으면 별표가 그대로 찍힌다
            for txt, b in _split_bold(t):
                r2 = par.add_run(txt)
                _apply_font(r2, size=10.5, bold=b)

    def formula(self, text):
        """수식 한 줄 — 고정폭으로 눈에 띄게."""
        par = self.doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.8)
        par.paragraph_format.space_before = Pt(4)
        par.paragraph_format.space_after = Pt(6)
        r = par.add_run(text)
        _apply_font(r, name="Consolas", size=10.5, color=DARK)
        return par

    def note(self, title, text, warn=False):
        """참고/주의 상자 — 좌측 색선 + 연한 배경."""
        par = self.doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.3)
        par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(8)
        color = "C77700" if warn else NAVY
        # 상자가 쪽 경계에서 갈라지면 색선만 남은 조각이 생긴다 — 통째로 넘긴다.
        par.paragraph_format.keep_together = True
        _border(par, color=color, size=18, sides=("left",))
        _shade(par, BOXBG_WARN if warn else BOXBG)
        r = par.add_run(title + "  ")
        _apply_font(r, size=10.5, bold=True, color=color)
        for txt, b in _split_bold(text):
            r2 = par.add_run(txt)
            _apply_font(r2, size=10.5, bold=b)
        return par

    def table(self, header, rows, widths=None, first_col_bold=False):
        """표.

        ⚠ 정렬 규칙(전 제품 공통 — 빼면 안 된다)
          · 모든 칸은 **세로 가운데**. python-docx 기본값은 위쪽 정렬이라
            반드시 `vertical_alignment` 를 직접 지정해야 한다. 칸마다 줄 수가
            다른 표(설명이 두 줄인 칸)에서 글자가 위로 붙어 어긋나 보인다.
          · **제목 행은 가로·세로 모두 가운데.** 문단 정렬은 스타일에서 물려받지
            않으므로 이것도 직접 준다.
        """
        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(header):
            hdr[i].text = ""
            par = hdr[i].paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER      # 제목 행 = 가로 가운데
            r = par.add_run(h)
            _apply_font(r, size=10, bold=True, color="FFFFFF")
            _shade(hdr[i], NAVY)
            hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                base = first_col_bold and i == 0
                # 칸 안에서도 `**굵게**` 를 쓸 수 있게 한다(표시가 없으면 종전과 똑같이 한 덩어리).
                for txt, bold in _split_bold(str(v)):
                    r = cells[i].paragraphs[0].add_run(txt)
                    _apply_font(r, size=10, bold=(base or bold))
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        # 표가 쪽 경계에서 쪼개지면 마지막 한 줄만 다음 장에 남아 보기 흉하다.
        # 행 자체의 분할을 막고(cantSplit), 마지막 행 앞까지 붙여 둔다(keep_with_next).
        for idx, row in enumerate(t.rows):
            trPr = row._tr.get_or_add_trPr()
            cant = OxmlElement("w:cantSplit")
            trPr.append(cant)
            if idx < len(t.rows) - 1:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        par.paragraph_format.keep_with_next = True
        # 긴 표가 쪽을 넘어갈 때 제목 행을 다시 보여 준다(넘어간 쪽에서 열 뜻을 알 수 있게).
        hdr_trPr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        hdr_trPr.append(th)
        self.p()
        return t

    # ⚠ 「표·그림 앞 문단을 본체와 붙이기」는 넣었다가 되돌렸다 (2026-08-13 실측).
    #   제목(keep_with_next) → 도입문 → 그림 → 그림설명 이 한 덩어리로 연결돼,
    #   그림이 쪽에 안 들어가면 제목까지 넷이 통째로 밀리면서 앞 쪽에 10~14cm 빈
    #   공간이 남았다. 묶음은 «갈라지면 안 되는 짝»(그림+설명, 표의 행)에만 걸고,
    #   나머지는 크기를 줄여서 해결한다. 다시 넣지 말 것.

    @staticmethod
    def _fit_width(path, want_w_cm):
        """원본 픽셀을 보고 폭을 정한다 — 폭·높이·확대배율 셋을 함께 만족시킨다.

        Pillow 가 없으면 요청한 폭을 그대로 쓴다(생성 자체는 막지 않는다).
        """
        try:
            from PIL import Image
            with Image.open(path) as im:
                pw, ph = im.size
        except Exception:
            return want_w_cm
        if not pw or not ph:
            return want_w_cm
        natural_w = pw / PX_PER_CM                       # 화면에서 보던 크기
        w = min(want_w_cm, natural_w * FIG_MAX_UPSCALE)   # 작은 조각을 억지로 늘리지 않는다
        if w * ph / pw > FIG_MAX_H_CM:                    # 높이가 상한을 넘으면 폭을 줄여 맞춘다
            w = FIG_MAX_H_CM * pw / ph
        return round(w, 2)

    def figure(self, path, caption, width_cm=FIG_W_CM):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 본문 줄간격(1.35)이 그림에도 걸리면 그림 아래에 빈 줄만큼 틈이 벌어진다.
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        # 그림과 그림 설명이 쪽 경계로 갈라지면 설명만 다음 장에 남는다.
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(path, width=Cm(self._fit_width(path, width_cm)))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.0
        r = cap.add_run(caption)
        _apply_font(r, size=9, color=CAPTION)
        cap.paragraph_format.space_after = Pt(10)

    def toc(self, entries):
        """목차 — Word 필드가 아니라 직접 쓴 목록(필드는 갱신 안 하면 쪽번호가 어긋난다)."""
        toc_head = self.doc.add_heading("목    차", level=1)
        # ⚠ 목차는 «한 쪽 안»에 들어가야 한다. 줄간격·여백을 조금만 키워도 마지막 한두 줄이
        #   다음 장으로 넘어가 거의 빈 쪽이 생긴다(항목이 늘 때마다 재발 — 2026-08-06 설치
        #   매뉴얼에서 항목 하나 늘었다고 3줄이 넘어갔다). 항목이 많으면 간격을 더 조인다.
        #   실측: 항목 40줄이면 여백 3pt 로는 3줄이 넘어간다. 32줄을 경계로 조인다.
        n_lines = len(entries) + sum(len(s) for _, s in entries)
        crowded = n_lines > 32
        gap_before = 0 if crowded else 3
        sub_size = 9.5 if crowded else 10
        if crowded:
            # 목차 제목의 위쪽 여백(기본 20pt)까지 줄여야 마지막 줄이 넘어가지 않는다.
            toc_head.paragraph_format.space_before = Pt(4)
            toc_head.paragraph_format.space_after = Pt(4)
        for title, subs in entries:
            par = self.doc.add_paragraph()
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(gap_before)
            par.paragraph_format.line_spacing = 1.0
            r = par.add_run(title)
            _apply_font(r, size=10.5, bold=True)
            for s in subs:
                sp = self.doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.line_spacing = 1.0
                sr = sp.add_run("  " + s)
                _apply_font(sr, size=sub_size)

    def save(self, path):
        self.doc.save(path)
        return path
